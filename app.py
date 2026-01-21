import io
import os
import pandas as pd
import streamlit as st
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.set_page_config(page_title="SEO Roadmap Builder (Prototype)", layout="wide")

DEFAULT_PLAYPACK_PATH = "SEO_PreMade_Plays_By_Client_Profile_v2.xlsx"

def read_default_playpack_bytes() -> bytes:
    if not os.path.exists(DEFAULT_PLAYPACK_PATH):
        raise FileNotFoundError(
            f"Default play pack not found at '{DEFAULT_PLAYPACK_PATH}'. "
            "Make sure the XLSX is in the repo root."
        )
    with open(DEFAULT_PLAYPACK_PATH, "rb") as f:
        return f.read()

def load_profiles(xlsx_bytes: bytes):
    wb = load_workbook(io.BytesIO(xlsx_bytes), data_only=True)
    profiles = {}
    for ws in wb.worksheets:
        header_row = None
        for r in range(1, 60):
            if ws.cell(row=r, column=1).value == "Strategy / Play":
                header_row = r
                break
        if not header_row:
            continue

        profile_name = ws.cell(row=1, column=2).value or ws.title
        profile_desc = ws.cell(row=2, column=2).value or ""

        rows = []
        r = header_row + 1
        while True:
            strat = ws.cell(row=r, column=1).value
            if not strat:
                break
            rows.append({
                "Strategy / Play": str(strat),
                "Priority Order": int(ws.cell(row=r, column=2).value or 999),
                "Month Allocation (1-6)": int(ws.cell(row=r, column=3).value or 6),
                "Allocated Effort (Minutes)": int(ws.cell(row=r, column=4).value or 0),
            })
            r += 1

        df = pd.DataFrame(rows).sort_values(["Priority Order", "Month Allocation (1-6)"]).reset_index(drop=True)
        profiles[ws.title] = {"name": profile_name, "desc": profile_desc, "df": df}
    return profiles

def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)

def add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_table(doc: Document, df: pd.DataFrame):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Light List Accent 1"
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = str(row[col])
    return table

def generate_docx(client_name: str, profile_label: str, profile_desc: str, plays_df: pd.DataFrame, topics_df: pd.DataFrame | None):
    doc = Document()

    title = doc.add_paragraph(f"{client_name} — SEO Action Plan (Prototype)")
    title.runs[0].font.size = Pt(20)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, "")

    add_heading(doc, "Overview", level=1)
    add_paragraph(doc, f"Client Profile: {profile_label}")
    if profile_desc:
        add_paragraph(doc, f"Profile Notes: {profile_desc}")
    add_paragraph(doc, "This document is a prototype skeleton generated from a standardized play-pack plus optional keyword/topic enrichment.")

    add_heading(doc, "Recommended Focus Areas", level=1)
    top_plays = plays_df.sort_values(["Priority Order"]).head(5)["Strategy / Play"].tolist()
    for p in top_plays:
        doc.add_paragraph(p, style="List Bullet")

    add_heading(doc, "6-Month Roadmap Timeline (Plays)", level=1)
    timeline = plays_df.sort_values(["Month Allocation (1-6)", "Priority Order"]).copy()
    timeline["Month"] = timeline["Month Allocation (1-6)"].apply(lambda x: f"Month {x}")
    timeline = timeline[["Month", "Priority Order", "Strategy / Play", "Allocated Effort (Minutes)"]]
    add_table(doc, timeline)

    add_heading(doc, "Keyword / Topic Enrichment (Optional)", level=1)
    if topics_df is None or topics_df.empty:
        add_paragraph(doc, "No keyword/topic upload provided. Roadmap uses generic placeholders where relevant.")
    else:
        add_paragraph(doc, "Uploaded topics/keywords were included to customize content-building plays.")
        show = topics_df.head(20).copy()
        add_table(doc, show)

    add_heading(doc, "Next Steps", level=1)
    add_paragraph(doc, "1) Review and adjust month allocation and priorities as needed.")
    add_paragraph(doc, "2) Add or refine specific tasks per play (URLs, pages, on-page edits, technical fixes).")
    add_paragraph(doc, "3) Confirm dependencies (access, dev support, approvals).")
    add_paragraph(doc, "4) Track execution and outcomes monthly (GSC, GA4, rank tracking).")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

st.title("SEO Roadmap Builder (Prototype)")
st.caption("Play packs are pre-loaded from the repo. Optional: override by uploading another play-pack XLSX.")

with st.sidebar:
    st.header("Inputs")
    client_name = st.text_input("Client name", value="Client Name")

    st.subheader("Play Pack")
    use_override = st.toggle("Override play pack (upload XLSX)", value=False)
    override_file = None
    if use_override:
        override_file = st.file_uploader("Upload Play Packs XLSX", type=["xlsx"], key="plays_override")
        st.caption("If OFF, the app uses the pre-loaded play pack in the repo.")

    topics_file = st.file_uploader("Keyword/Topic CSV (optional)", type=["csv"], help="Optional enrichment file.", key="topics")

    st.divider()
    st.header("Audit overrides (prototype)")
    tech_blocker = st.checkbox("Tech blocker present (indexing/crawl)", value=False)
    local_priority = st.checkbox("Local/GBP priority", value=True)
    content_gap = st.checkbox("Content gaps present", value=True)

# Load play pack (default or override)
try:
    if use_override and override_file is not None:
        xlsx_bytes = override_file.read()
        playpack_label = "Uploaded override"
    else:
        xlsx_bytes = read_default_playpack_bytes()
        playpack_label = "Pre-loaded default"
except Exception as e:
    st.error(f"Could not load play pack: {e}")
    st.stop()

profiles = load_profiles(xlsx_bytes)
profile_keys = list(profiles.keys())
if not profile_keys:
    st.error("No profiles found in the play pack XLSX (missing Strategy table headers).")
    st.stop()

st.info(f"Using play pack: **{playpack_label}**")

profile_choice = st.selectbox(
    "Select client profile",
    profile_keys,
    format_func=lambda k: f"{k} — {profiles[k]['desc']}"
)

df = profiles[profile_choice]["df"].copy()

# Prototype overrides
if tech_blocker:
    for strat in ["Indexability Assurance", "Technical Integrity & Error Resolution"]:
        mask = df["Strategy / Play"].eq(strat)
        if mask.any():
            df.loc[mask, "Month Allocation (1-6)"] = 1
            df.loc[mask, "Priority Order"] = 0
if local_priority:
    mask = df["Strategy / Play"].eq("Local SEO & GBP Optimization")
    if mask.any():
        df.loc[mask, "Month Allocation (1-6)"] = 1
if not content_gap:
    mask = df["Strategy / Play"].eq("Horizontal Content Expansion")
    if mask.any():
        df.loc[mask, "Month Allocation (1-6)"] = 5

df = df.sort_values(["Month Allocation (1-6)", "Priority Order"]).reset_index(drop=True)

st.subheader("Play Pack (editable)")
edited = st.data_editor(
    df,
    use_container_width=True,
    num_rows="dynamic",
    column_config={
        "Priority Order": st.column_config.NumberColumn(min_value=0, step=1),
        "Month Allocation (1-6)": st.column_config.NumberColumn(min_value=1, max_value=6, step=1),
        "Allocated Effort (Minutes)": st.column_config.NumberColumn(min_value=0, step=5),
    },
)

total_minutes = int(edited["Allocated Effort (Minutes)"].sum())
st.metric("Total effort", f"{total_minutes} mins", f"{round(total_minutes/60, 2)} hrs")

topics_df = None
if topics_file is not None:
    topics_df = pd.read_csv(topics_file)
    st.subheader("Keyword/topic preview")
    st.dataframe(topics_df.head(25), use_container_width=True)

if st.button("Generate DOCX skeleton"):
    docx_buf = generate_docx(
        client_name=client_name.strip() or "Client Name",
        profile_label=profiles[profile_choice]["name"] or profile_choice,
        profile_desc=profiles[profile_choice]["desc"],
        plays_df=edited,
        topics_df=topics_df
    )
    st.download_button(
        "Download DOCX",
        data=docx_buf,
        file_name=f"{client_name or 'client'}_SEO_Action_Plan_Prototype.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
